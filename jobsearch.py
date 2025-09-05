import os
import requests
import json
import time
from datetime import datetime
from docx import Document
import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font

# -------------------------
# Configuration
# -------------------------
USER_CONFIG = {
    "cv": "C:/Users/harou/Documents/CV/Bacha Haroun Errachid CV.pdf",
    "name": "Haroun Bacha",
    "email": "harounbacha2005@gmail.com",
    "phone": "+213656536613",
    "degree": "bachelor's",
    "location": "Batna, Algeria",
     "queries": [
        "IT",
        "Ausbildung",
        "software developer",
        "intern",
        "web development",
        "frontend",
        "backend",
        "fullstack",
        "programming",
        "computer science",
        "apprenticeship",
        "junior developer",
        "data analyst"
    ],
     "locations": [
        "Germany",
        "Algeria",
        "France",
        "Uunited states",
        "Canada",
        "United Kingdom",
        "Netherlands",
        "Remote"
    ],
    "JSEARCH_API_KEY": "e5bd20b0f6mshf396f33cfefb4b4p1361f4jsn2bdd3fa427a0",
    "JSEARCH_HOST": "jsearch.p.rapidapi.com",
    "jsearch_num_pages": 3,
    "JOBBOERSE_API_KEY": os.getenv("JOBBOERSE_API_KEY", ""),
    "output_dir": "./application_bot_outputs",
}

os.makedirs(USER_CONFIG["output_dir"], exist_ok=True)

# -------------------------
# Utilities
# -------------------------
HEADERS_JS = lambda key: {
    'x-rapidapi-key': key,
    'x-rapidapi-host': USER_CONFIG['JSEARCH_HOST']
}

def safe_get(url, headers=None, params=None, retries=2, timeout=12):
    for attempt in range(retries):
        try:
            r = requests.get(url, headers=headers, params=params, timeout=timeout)
            r.raise_for_status()
            return r
        except requests.RequestException as e:
            print(f"Request error: {e} (attempt {attempt+1}/{retries})")
            time.sleep(1 + attempt * 2)
    return None

# -------------------------
# JSearch Integration
# -------------------------

def search_jsearch(query, location, api_key, num_pages=1):
    if not api_key:
        print("⚠️ JSearch API key not set. Skipping JSearch queries.")
        return []

    results = []
    base = f"https://{USER_CONFIG['JSEARCH_HOST']}/search"

    for page in range(1, num_pages + 1):
        params = {'query': f"{query} {location}", 'page': page, 'num_pages': 1}
        print(f"JSearch: querying '{params['query']}' (page {page})")
        resp = safe_get(base, headers=HEADERS_JS(api_key), params=params)
        if not resp:
            continue
        try:
            data = resp.json()
        except json.JSONDecodeError:
            print("Failed to decode JSearch JSON")
            continue

        items = data.get('data') or data.get('jobs') or []
        for it in items:
            results.append({
                'title': it.get('job_title') or it.get('title') or '',
                'company': it.get('employer_name') or it.get('company_name') or '',
                'location': it.get('job_city') or it.get('location') or location,
                'requirements': it.get('snippet') or it.get('job_description') or '',
                'date_posted': it.get('job_posted_at_datetime_utc') or '',
                'link': it.get('job_apply_link') or it.get('url') or '',
                'source': 'JSearch'
            })
        time.sleep(0.6)
    print(f"JSearch: found {len(results)} items for '{query} {location}'")
    return results

# -------------------------
# Jobbörse Placeholder
# -------------------------

def search_jobboerse(query, location, api_key, max_pages=1):
    if not api_key:
        print("Jobbörse API key not set. Skipping Jobbörse search.")
        return []
    print("Jobbörse search placeholder — implement API call here.")
    return []

# -------------------------
# Excel Export
# -------------------------
EXCEL_COLUMNS = ['Source','Title','Company','Location','Requirements/Description','Date Posted','Link','Status','Notes']

def save_to_excel(items, filename):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Opportunities'
    for idx, col in enumerate(EXCEL_COLUMNS, start=1):
        ws.cell(row=1, column=idx, value=col).font = Font(bold=True)

    for row_idx, item in enumerate(items, start=2):
        ws.cell(row=row_idx, column=1, value=item.get('source'))
        ws.cell(row=row_idx, column=2, value=item.get('title'))
        ws.cell(row=row_idx, column=3, value=item.get('company'))
        ws.cell(row=row_idx, column=4, value=item.get('location'))
        ws.cell(row=row_idx, column=5, value=item.get('requirements'))
        ws.cell(row=row_idx, column=6, value=item.get('date_posted'))
        ws.cell(row=row_idx, column=7, value=item.get('link'))
        ws.cell(row=row_idx, column=8, value='Pending')
        ws.cell(row=row_idx, column=9, value='')

    for i, column in enumerate(ws.columns, start=1):
        max_length = max(len(str(cell.value or '')) for cell in column)
        ws.column_dimensions[get_column_letter(i)].width = min(max(15, max_length+2), 60)

    wb.save(filename)
    print(f"Saved Excel to: {filename}")

# -------------------------
# Email & Motivation Letters
# -------------------------

def generate_email_text(item, user_config):
    recipient = item.get('company') or 'Hiring Team'
    title = item.get('title') or 'the position'
    
    text = f"Subject: Application for {title}\n\n"
    text += f"Dear {recipient},\n\n"
    text += (
        f"My name is {user_config['name']}, and I am writing to express my interest in the {title} role at your company. "
        f"I hold a {user_config['degree']} and am eager to contribute my skills and enthusiasm to your team.\n\n"
        f"Attached is my CV for your review. I would be delighted to discuss how I can add value to your organization. "
        f"You can reach me at {user_config['email']} or {user_config['phone']}.\n\n"
        f"Thank you for considering my application. I look forward to the possibility of contributing to {recipient}.\n\n"
        f"Sincerely,\n"
        f"{user_config['name']}\n"
        f"{user_config['location']}"
        f"{user_config['CV']}"
    )
    
    return text


def save_email(item, user_config, out_dir, filename):
    path = os.path.join(out_dir, filename)
    with open(path,'w',encoding='utf-8') as f:
        f.write(generate_email_text(item,user_config))
    print(f"Saved email draft: {path}")

def generate_motivation_doc(item, user_config, out_dir, filename):
    doc = Document()
    
    # Heading
    doc.add_heading('Motivation Letter', level=1)
    
    # Greeting
    recipient = item.get('company') or 'Selection Committee'
    doc.add_paragraph(f"Dear {recipient},")
    
    # Introduction
    doc.add_paragraph(
        f"My name is {user_config['name']}, and I hold a {user_config['degree']} from Algeria. "
        f"I am very interested in the position of '{item.get('title')}' because "
        f"{item.get('requirements') or 'it aligns perfectly with my skills, experience, and career goals.'}"
    )
    
    # Motivation and skills
    doc.add_paragraph(
        "I am highly motivated to grow professionally and contribute meaningfully to your team. "
        "I am a quick learner, adaptable, and eager to apply my knowledge to help achieve organizational goals."
    )
    
    # Closing
    doc.add_paragraph(
        f"Thank you for considering my application. I would be delighted to discuss how I can contribute to {recipient}.\n\n"
        f"Sincerely,\n"
        f"{user_config['name']}\n"
        f"{user_config['email']} | {user_config['phone']}\n"
        f"{user_config.get('location', '')}"
    )
    
    # Save document
    path = os.path.join(out_dir, filename)
    doc.save(path)
    print(f"Saved motivation letter: {path}")

# -------------------------
# Main Bot
# -------------------------

def build_tasks(user_config):
    return [{'query':q,'location':loc} for q in user_config['queries'] for loc in user_config['locations']]

def run_bot(user_config):
    all_results = []
    tasks = build_tasks(user_config)
    jsearch_key = user_config.get('JSEARCH_API_KEY')

    for t in tasks:
        all_results.extend(search_jsearch(t['query'], t['location'], jsearch_key, user_config.get('jsearch_num_pages',1)))

    jobboerse_key = user_config.get('JOBBOERSE_API_KEY')
    if jobboerse_key:
        for t in tasks:
            all_results.extend(search_jobboerse(t['query'], t['location'], jobboerse_key))

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    excel_path = os.path.join(user_config['output_dir'], f'opportunities_{timestamp}.xlsx')
    save_to_excel(all_results, excel_path)

    for idx, item in enumerate(all_results, start=1):
        safe_title = ''.join(c if c.isalnum() else '_' for c in (item.get('title') or 'opportunity'))
        email_file = f'{idx}_{safe_title[:40]}_email.txt'
        save_email(item, user_config, user_config['output_dir'], email_file)        
        
        mot_file = f'{idx}_{safe_title[:40]}_motivation.docx'       
        generate_motivation_doc(item, user_config, user_config['output_dir'], mot_file)     

if __name__ == '__main__':
    run_bot(USER_CONFIG)
